from flask import Flask, render_template, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import os
import whisper
import torch
import librosa
import soundfile as sf
from pydub import AudioSegment
import tempfile
import openai
from textblob import TextBlob
import nltk
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
import numpy as np
import pandas as pd
import re
from collections import Counter
import json
import uuid
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from docx.shared import Inches
import io
from call_service import CallService
from recording_service import RecordingService
from webhook_service import WebhookService
from dotenv import load_dotenv
from aws_transcribe_streaming import AWSTranscribeStreamingService

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*")

# Initialize services
call_service = CallService()
recording_service = RecordingService()
webhook_service = WebhookService(call_service, recording_service, socketio)

# Initialize AWS Transcribe Streaming service after other services
aws_transcribe_service = AWSTranscribeStreamingService(socketio)

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
# Email configuration - Updated with environment variables and fallback
EMAIL_HOST = os.getenv('EMAIL_HOST', 'smtp.gmail.com')
EMAIL_PORT = int(os.getenv('EMAIL_PORT', 587))
EMAIL_USER = os.getenv('EMAIL_USER', None)  # Will be None if not configured
EMAIL_PASSWORD = os.getenv('EMAIL_PASSWORD', None)  # Will be None if not configured
EMAIL_ENABLED = EMAIL_USER is not None and EMAIL_PASSWORD is not None

# Alternative email providers configuration
EMAIL_PROVIDERS = {
    'gmail': {
        'host': 'smtp.gmail.com',
        'port': 587,
        'tls': True
    },
    'outlook': {
        'host': 'smtp-mail.outlook.com',
        'port': 587,
        'tls': True
    },
    'yahoo': {
        'host': 'smtp.mail.yahoo.com',
        'port': 587,
        'tls': True
    }
}

# Create upload directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Load Whisper model
print("Loading Whisper model...")
try:
    whisper_model = whisper.load_model("base")
    print("Whisper model loaded successfully!")
except Exception as e:
    print(f"Error loading Whisper model: {e}")
    whisper_model = None
print("Whisper model loaded successfully!")

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('vader_lexicon', quiet=True)
    nltk.download('wordnet', quiet=True)
except:
    print("Warning: Some NLTK data downloads failed")

@app.route('/')
def index():
    """Main page route"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_audio():
    """Handle audio file upload"""
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        file = request.files['audio']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Extended audio format validation
        allowed_extensions = {
            'wav', 'mp3', 'mp4', 'mpeg', 'mpga', 'm4a', 'webm', 
            'ogg', 'flac', 'aac', 'wma', '3gp', 'amr', 'aiff'
        }
        
        if not ('.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions):
            return jsonify({'error': f'Invalid file type. Supported formats: {", ".join(allowed_extensions)}'}), 400
        
        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_extension = file.filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{timestamp}_{uuid.uuid4().hex[:8]}.{file_extension}"
        
        # Ensure upload directory exists
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Save file
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)

        if whisper_model is None:
            return jsonify({'error': 'Whisper model not available. Please reinstall whisper.'}), 500
        
        # Get file info
        file_size = os.path.getsize(file_path)
        
        return jsonify({
            'message': 'File uploaded successfully',
            'filename': unique_filename,
            'original_filename': file.filename,
            'file_size': file_size,
            'file_path': file_path
        })
        
    except Exception as e:
        print(f"Upload error: {str(e)}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """Convert audio to text using OpenAI Whisper"""
    try:
        data = request.get_json()
        if not data or 'filename' not in data:
            return jsonify({'error': 'No filename provided'}), 400
        
        filename = data['filename']
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Convert audio to compatible format if needed
        audio_path = preprocess_audio(file_path)
        
        # Transcribe using Whisper with word-level timestamps
        print(f"Transcribing audio file: {filename}")
        result = whisper_model.transcribe(audio_path, word_timestamps=True)
        
        # Extract transcription and segments with word-level timestamps
        transcription = result['text']
        segments = []
        
        for segment in result['segments']:
            segment_data = {
                'start': segment['start'],
                'end': segment['end'],
                'text': segment['text'],
                'words': []
            }
            
            # Add word-level timestamps if available
            if 'words' in segment:
                for word in segment['words']:
                    segment_data['words'].append({
                        'start': word['start'],
                        'end': word['end'],
                        'word': word['word']
                    })
            
            segments.append(segment_data)
        
        # Clean up temporary file if created
        if audio_path != file_path:
            os.remove(audio_path)
        
        return jsonify({
            'transcription': transcription,
            'segments': segments,
            'language': result.get('language', 'unknown'),
            'status': 'success'
        })
        
    except Exception as e:
        print(f"Transcription error: {str(e)}")
        return jsonify({'error': f'Transcription failed: {str(e)}'}), 500

def preprocess_audio(file_path):
    """Convert audio file to format compatible with Whisper"""
    try:
        # Get file extension
        file_extension = file_path.split('.')[-1].lower()
        
        # If already in compatible format, return as is
        if file_extension in ['wav', 'mp3', 'flac', 'm4a']:
            return file_path
        
        # Convert to WAV format
        audio = AudioSegment.from_file(file_path)
        
        # Create temporary WAV file
        temp_path = tempfile.mktemp(suffix='.wav')
        audio.export(temp_path, format='wav')
        
        return temp_path
        
    except Exception as e:
        print(f"Audio preprocessing error: {str(e)}")
        return file_path

@app.route('/analyze', methods=['POST'])
def analyze_text():
    """Analyze transcribed text for insights"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided for analysis'}), 400
        
        text = data['text']
        if not text.strip():
            return jsonify({'error': 'Empty text provided'}), 400
        
        # Perform comprehensive analysis
        analysis_results = {
            'sentiment_analysis': analyze_sentiment(text),
            'emotion_analysis': analyze_emotions(text),
            'key_topics': extract_key_topics(text),
            'insights': extract_insights(text),
            'summary': generate_summary(text),
            'suggestions': generate_suggestions(text),
            'statistics': get_text_statistics(text)
        }
        
        return jsonify({
            'status': 'success',
            'analysis': analysis_results
        })
        
    except Exception as e:
        print(f"Analysis error: {str(e)}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

def analyze_sentiment(text):
    """Analyze sentiment using TextBlob"""
    try:
        blob = TextBlob(text)
        polarity = blob.sentiment.polarity
        subjectivity = blob.sentiment.subjectivity
        
        # Determine sentiment category
        if polarity > 0.1:
            sentiment = "Positive"
        elif polarity < -0.1:
            sentiment = "Negative"
        else:
            sentiment = "Neutral"
        
        return {
            'sentiment': sentiment,
            'polarity': round(polarity, 3),
            'subjectivity': round(subjectivity, 3),
            'confidence': round(abs(polarity), 3)
        }
    except Exception as e:
        return {'error': str(e)}

def analyze_emotions(text):
    """Analyze emotions using keyword-based approach"""
    try:
        emotion_keywords = {
            'joy': ['happy', 'joy', 'excited', 'pleased', 'delighted', 'cheerful', 'glad', 'satisfied'],
            'sadness': ['sad', 'unhappy', 'depressed', 'disappointed', 'grief', 'sorrow', 'upset'],
            'anger': ['angry', 'mad', 'furious', 'irritated', 'annoyed', 'frustrated', 'rage'],
            'fear': ['afraid', 'scared', 'worried', 'anxious', 'nervous', 'terrified', 'panic'],
            'surprise': ['surprised', 'shocked', 'amazed', 'astonished', 'stunned', 'bewildered'],
            'disgust': ['disgusted', 'revolted', 'repulsed', 'sickened', 'appalled']
        }
        
        text_lower = text.lower()
        emotion_scores = {}
        
        for emotion, keywords in emotion_keywords.items():
            score = sum(text_lower.count(keyword) for keyword in keywords)
            emotion_scores[emotion] = score
        
        # Normalize scores
        total_words = len(text.split())
        if total_words > 0:
            for emotion in emotion_scores:
                emotion_scores[emotion] = round(emotion_scores[emotion] / total_words * 100, 2)
        
        # Find dominant emotion
        dominant_emotion = max(emotion_scores, key=emotion_scores.get) if max(emotion_scores.values()) > 0 else 'neutral'
        
        return {
            'dominant_emotion': dominant_emotion,
            'emotion_scores': emotion_scores,
            'emotional_intensity': round(max(emotion_scores.values()), 2)
        }
    except Exception as e:
        return {'error': str(e)}

def extract_key_topics(text):
    """Extract key topics using improved keyword extraction"""
    try:
        # Get keywords first
        keywords = get_top_keywords(text)
        
        # Try TF-IDF if available
        try:
            import nltk
            from sklearn.feature_extraction.text import TfidfVectorizer
            
            sentences = nltk.sent_tokenize(text)
            if len(sentences) >= 2:
                # TF-IDF vectorization
                vectorizer = TfidfVectorizer(
                    max_features=50,
                    stop_words='english',
                    ngram_range=(1, 2),
                    min_df=1
                )
                
                tfidf_matrix = vectorizer.fit_transform(sentences)
                feature_names = vectorizer.get_feature_names_out()
                
                # Get top keywords
                import numpy as np
                mean_scores = np.mean(tfidf_matrix.toarray(), axis=0)
                top_indices = mean_scores.argsort()[-10:][::-1]
                
                topics = []
                for idx in top_indices:
                    if mean_scores[idx] > 0:
                        topics.append({
                            'topic': feature_names[idx],
                            'relevance': round(mean_scores[idx], 3)
                        })
                
                return {
                    'topics': topics,
                    'keywords': keywords
                }
        except:
            pass
        
        # Fallback: Use simple keyword-based topics
        text_lower = text.lower()
        common_topics = {
            'customer service': ['customer', 'service', 'help', 'support', 'assist'],
            'communication': ['call', 'phone', 'speak', 'talk', 'conversation'],
            'questions': ['question', 'ask', 'wonder', 'inquire', 'what', 'how', 'why'],
            'assistance': ['help', 'assist', 'support', 'aid', 'guide'],
            'greeting': ['hello', 'hi', 'thank', 'please', 'welcome']
        }
        
        topics = []
        for topic, topic_words in common_topics.items():
            score = sum(text_lower.count(word) for word in topic_words)
            if score > 0:
                topics.append({
                    'topic': topic,
                    'relevance': round(score / len(text.split()) * 100, 3)
                })
        
        # Sort by relevance
        topics.sort(key=lambda x: x['relevance'], reverse=True)
        
        return {
            'topics': topics[:10],  # Top 10 topics
            'keywords': keywords
        }
        
    except Exception as e:
        print(f"Topics error: {e}")
        return {
            'topics': [],
            'keywords': get_top_keywords(text)
        }
def get_top_keywords(text):
    """Extract top keywords from text"""
    try:
        # Remove punctuation and convert to lowercase
        text_clean = re.sub(r'[^\w\s]', '', text.lower())
        words = text_clean.split()
        
        # Filter out common stop words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them'}
        
        filtered_words = [word for word in words if word not in stop_words and len(word) > 2]
        
        # Count word frequencies
        word_freq = Counter(filtered_words)
        
        # Get top 10 keywords
        top_keywords = []
        for word, freq in word_freq.most_common(10):
            top_keywords.append({
                'keyword': word,
                'frequency': freq
            })
        
        return top_keywords
    except Exception as e:
        return []

def extract_insights(text):
    """Extract insights from text"""
    try:
        insights = []
        
        # Basic text analysis
        words = text.split()
        word_count = len(words)
        
        # Use basic sentence splitting as fallback
        try:
            import nltk
            sentences = nltk.sent_tokenize(text)
        except:
            sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        sentence_count = len(sentences)
        
        # Text length insights
        if word_count > 500:
            insights.append("This is a lengthy text that covers substantial content.")
        elif word_count < 50:
            insights.append("This is a brief text with concise information.")
        else:
            insights.append("This text has a moderate length with balanced content.")
        
        # Complexity insights
        if sentence_count > 0:
            avg_sentence_length = word_count / sentence_count
            if avg_sentence_length > 20:
                insights.append("The text contains complex, detailed sentences.")
            elif avg_sentence_length < 10:
                insights.append("The text uses simple, concise sentences.")
            else:
                insights.append("The text has well-balanced sentence structure.")
        
        # Content type insights
        text_lower = text.lower()
        if 'customer service' in text_lower or 'help' in text_lower or 'support' in text_lower:
            insights.append("This appears to be a customer service interaction.")
        
        if 'hello' in text_lower or 'hi' in text_lower or 'thank you' in text_lower:
            insights.append("The text contains polite greetings and courteous language.")
        
        # Question insights
        question_count = text.count('?')
        if question_count > 3:
            insights.append("The text contains many questions, suggesting inquiry or discussion.")
        elif question_count > 0:
            insights.append("The text includes questions for engagement.")
        
        # Exclamation insights
        exclamation_count = text.count('!')
        if exclamation_count > 2:
            insights.append("The text shows strong emotions or emphasis.")
        
        # Personal pronouns
        personal_pronouns = [' i ', ' me ', ' my ', ' mine ', ' myself ']
        personal_count = sum(text_lower.count(pronoun) for pronoun in personal_pronouns)
        if personal_count > 3:
            insights.append("The text is highly personal and subjective.")
        
        # Professional language
        professional_words = ['service', 'help', 'assist', 'support', 'thank', 'please']
        professional_count = sum(text_lower.count(word) for word in professional_words)
        if professional_count > 2:
            insights.append("The text uses professional and courteous language.")
        
        return insights if insights else ["This text contains conversational content."]
        
    except Exception as e:
        print(f"Insights error: {e}")
        # Fallback insights
        word_count = len(text.split())
        if word_count > 100:
            return ["This is a substantial piece of text with multiple topics."]
        elif word_count > 50:
            return ["This is a moderate-length text with clear content."]
        else:
            return ["This is a brief text with concise information."]
def generate_summary(text):
    """Generate a summary of the text"""
    try:
        # Use basic sentence splitting as fallback
        try:
            import nltk
            sentences = nltk.sent_tokenize(text)
        except:
            sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
        
        if len(sentences) <= 3:
            return text
        
        # Simple extractive summarization
        words = text.lower().split()
        word_freq = {}
        for word in words:
            if len(word) > 3:  # Skip short words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        sentence_scores = {}
        for sentence in sentences:
            sentence_words = sentence.lower().split()
            score = sum(word_freq.get(word, 0) for word in sentence_words)
            sentence_scores[sentence] = score
        
        # Get top 3 sentences
        if sentence_scores:
            top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:3]
            summary = ' '.join([sentence[0] for sentence in top_sentences])
            return summary if summary else text[:200] + "..."
        else:
            return text[:200] + "..." if len(text) > 200 else text
            
    except Exception as e:
        print(f"Summary error: {e}")
        return text[:200] + "..." if len(text) > 200 else text
def generate_suggestions(text):
    """Generate suggestions based on text analysis"""
    try:
        suggestions = []
        
        # Analyze sentiment for suggestions
        try:
            sentiment_result = analyze_sentiment(text)
            if isinstance(sentiment_result, dict) and 'error' not in sentiment_result:
                sentiment = sentiment_result.get('sentiment', 'Neutral')
                if sentiment == 'Negative':
                    suggestions.append("Consider addressing the negative aspects mentioned to improve overall tone.")
                elif sentiment == 'Positive':
                    suggestions.append("Great positive tone! Consider leveraging this enthusiasm in future communications.")
                else:
                    suggestions.append("The neutral tone is balanced. Consider adding more emotional engagement.")
        except:
            suggestions.append("Consider reviewing the overall tone of the content.")
        
        # Length-based suggestions
        words = text.split()
        word_count = len(words)
        if word_count > 1000:
            suggestions.append("Consider breaking down this lengthy content into smaller, digestible sections.")
        elif word_count < 100:
            suggestions.append("Consider expanding on key points to provide more comprehensive information.")
        else:
            suggestions.append("The content length is appropriate for the topic.")
        
        # Clarity suggestions
        try:
            import nltk
            sentences = nltk.sent_tokenize(text)
        except:
            sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        sentence_count = len(sentences)
        if sentence_count > 0:
            avg_sentence_length = word_count / sentence_count
            if avg_sentence_length > 25:
                suggestions.append("Consider using shorter sentences to improve readability.")
            elif avg_sentence_length < 8:
                suggestions.append("Consider combining some short sentences for better flow.")
        
        # Engagement suggestions
        question_count = text.count('?')
        if question_count == 0:
            suggestions.append("Consider adding questions to increase engagement with your audience.")
        
        # Action-oriented suggestions
        action_words = ['should', 'must', 'need', 'recommend', 'suggest', 'propose']
        action_count = sum(text.lower().count(word) for word in action_words)
        if action_count == 0:
            suggestions.append("Consider adding actionable recommendations or next steps.")
        
        # Customer service specific suggestions
        text_lower = text.lower()
        if 'customer service' in text_lower or 'support' in text_lower:
            suggestions.append("For customer service interactions, ensure clear resolution steps are provided.")
        
        return suggestions if suggestions else ["The text appears well-structured overall."]
        
    except Exception as e:
        print(f"Suggestions error: {e}")
        return ["Consider reviewing the content for clarity and engagement."]
def get_text_statistics(text):
    """Get basic text statistics"""
    try:
        # Use basic sentence splitting as fallback
        try:
            import nltk
            sentences = nltk.sent_tokenize(text)
        except:
            sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        words = text.split()
        
        # Calculate statistics
        word_count = len(words)
        sentence_count = len(sentences)
        character_count = len(text)
        
        # Calculate averages
        avg_sentence_length = round(word_count / sentence_count, 2) if sentence_count > 0 else 0
        avg_word_length = round(sum(len(word) for word in words) / word_count, 2) if word_count > 0 else 0
        reading_time_minutes = round(word_count / 200, 1)  # Assuming 200 words per minute
        
        return {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'character_count': character_count,
            'avg_sentence_length': avg_sentence_length,
            'avg_word_length': avg_word_length,
            'reading_time_minutes': reading_time_minutes
        }
    except Exception as e:
        print(f"Statistics error: {e}")
        # Fallback statistics
        words = text.split()
        return {
            'word_count': len(words),
            'sentence_count': text.count('.') + text.count('!') + text.count('?'),
            'character_count': len(text),
            'avg_sentence_length': 0,
            'avg_word_length': 0,
            'reading_time_minutes': round(len(words) / 200, 1)
        }
@app.route('/download/json', methods=['POST'])
def download_json():
    """Generate JSON report for download"""
    try:
        data = request.get_json()
        if not data or 'report_data' not in data:
            return jsonify({'error': 'Report data required'}), 400
        
        report_data = data['report_data']
        
        # Create comprehensive JSON report
        json_report = {
            'report_info': {
                'generated_at': datetime.now().isoformat(),
                'platform': 'Vocalytics',
                'version': '1.0'
            },
            'transcription': report_data.get('transcription', ''),
            'language': report_data.get('language', 'unknown'),
            'analysis': {
                'sentiment_analysis': report_data.get('sentiment', {}),
                'emotion_analysis': report_data.get('emotions', {}),
                'key_topics': report_data.get('topics', {}),
                'statistics': report_data.get('statistics', {}),
                'summary': report_data.get('summary', ''),
                'insights': report_data.get('insights', []),
                'suggestions': report_data.get('suggestions', [])
            }
        }
        
        return jsonify({
            'filename': f'vocalytics_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
            'content': json.dumps(json_report, indent=2),
            'mime_type': 'application/json'
        })
        
    except Exception as e:
        print(f"JSON download error: {str(e)}")
        return jsonify({'error': f'Failed to generate JSON report: {str(e)}'}), 500

@app.route('/download/word', methods=['POST'])
def download_word():
    """Generate Word document report for download"""
    try:
        data = request.get_json()
        if not data or 'report_data' not in data:
            return jsonify({'error': 'Report data required'}), 400
        
        report_data = data['report_data']
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Vocalytics Audio Analysis Report', 0)
        title.alignment = 1  # Center alignment
        
        # Add generated date
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph("=" * 50)
        
        # Add transcription section
        doc.add_heading('Transcription', level=1)
        doc.add_paragraph(report_data.get('transcription', 'N/A'))
        
        # Add language info
        if report_data.get('language'):
            doc.add_paragraph(f"Detected Language: {report_data.get('language', 'Unknown')}")
        
        # Add sentiment analysis
        doc.add_heading('Sentiment Analysis', level=1)
        sentiment = report_data.get('sentiment', {})
        doc.add_paragraph(f"Overall Sentiment: {sentiment.get('sentiment', 'N/A')}")
        doc.add_paragraph(f"Polarity: {sentiment.get('polarity', 'N/A')}")
        doc.add_paragraph(f"Subjectivity: {sentiment.get('subjectivity', 'N/A')}")
        doc.add_paragraph(f"Confidence: {sentiment.get('confidence', 'N/A')}")
        
        # Add emotion analysis
        doc.add_heading('Emotion Analysis', level=1)
        emotions = report_data.get('emotions', {})
        doc.add_paragraph(f"Dominant Emotion: {emotions.get('dominant_emotion', 'N/A')}")
        doc.add_paragraph(f"Emotional Intensity: {emotions.get('emotional_intensity', 'N/A')}%")
        
        if emotions.get('emotion_scores'):
            doc.add_paragraph("Emotion Breakdown:")
            for emotion, score in emotions['emotion_scores'].items():
                doc.add_paragraph(f"  • {emotion.title()}: {score}%")
        
        # Add statistics
        doc.add_heading('Text Statistics', level=1)
        stats = report_data.get('statistics', {})
        doc.add_paragraph(f"Word Count: {stats.get('word_count', 'N/A')}")
        doc.add_paragraph(f"Sentence Count: {stats.get('sentence_count', 'N/A')}")
        doc.add_paragraph(f"Average Sentence Length: {stats.get('avg_sentence_length', 'N/A')}")
        doc.add_paragraph(f"Reading Time: {stats.get('reading_time_minutes', 'N/A')} minutes")
        
        # Add summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(report_data.get('summary', 'N/A'))
        
        # Add insights
        doc.add_heading('Insights', level=1)
        insights = report_data.get('insights', [])
        if insights:
            for insight in insights:
                doc.add_paragraph(f"• {insight}")
        else:
            doc.add_paragraph('No insights available')
        
        # Add suggestions
        doc.add_heading('Suggestions', level=1)
        suggestions = report_data.get('suggestions', [])
        if suggestions:
            for suggestion in suggestions:
                doc.add_paragraph(f"• {suggestion}")
        else:
            doc.add_paragraph('No suggestions available')
        
        # Save to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Convert to base64 for JSON response
        import base64
        doc_base64 = base64.b64encode(doc_io.getvalue()).decode('utf-8')
        
        return jsonify({
            'filename': f'vocalytics_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            'content': doc_base64,
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        })
        
    except Exception as e:
        print(f"Word download error: {str(e)}")
        return jsonify({'error': f'Failed to generate Word report: {str(e)}'}), 500

@app.route('/send-email', methods=['POST'])
def send_email():
    """Send analysis report via email with improved error handling"""
    try:
        data = request.get_json()
        if not data or 'email' not in data or 'report_data' not in data:
            return jsonify({'error': 'Email and report data required'}), 400
        
        recipient_email = data['email']
        report_data = data['report_data']
        
        # Check if email is configured
        if not EMAIL_ENABLED:
            # Return demo response when email is not configured
            return jsonify({
                'success': True,
                'message': 'Email functionality is in demo mode. In production, configure EMAIL_USER and EMAIL_PASSWORD environment variables.',
                'demo_mode': True,
                'recipient': recipient_email,
                'report_summary': {
                    'transcription_length': len(report_data.get('transcription', '')),
                    'sentiment': report_data.get('sentiment', {}).get('sentiment', 'N/A'),
                    'dominant_emotion': report_data.get('emotions', {}).get('dominant_emotion', 'N/A')
                }
            })
        
        # Validate email address format
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, recipient_email):
            return jsonify({'error': 'Invalid email address format'}), 400
        
        # Create email message
        msg = MIMEMultipart()
        msg['From'] = EMAIL_USER
        msg['To'] = recipient_email
        msg['Subject'] = 'Vocalytics - Audio Analysis Report'
        
        # Create email body
        body = f"""
        Dear User,

        Thank you for using Vocalytics! Please find your audio analysis report below:

        TRANSCRIPTION:
        {report_data.get('transcription', 'N/A')}

        SENTIMENT ANALYSIS:
        - Overall Sentiment: {report_data.get('sentiment', {}).get('sentiment', 'N/A')}
        - Polarity: {report_data.get('sentiment', {}).get('polarity', 'N/A')}
        - Subjectivity: {report_data.get('sentiment', {}).get('subjectivity', 'N/A')}

        EMOTION ANALYSIS:
        - Dominant Emotion: {report_data.get('emotions', {}).get('dominant_emotion', 'N/A')}
        - Emotional Intensity: {report_data.get('emotions', {}).get('emotional_intensity', 'N/A')}%

        SUMMARY:
        {report_data.get('summary', 'N/A')}

        INSIGHTS:
        {chr(10).join(report_data.get('insights', []))}

        SUGGESTIONS:
        {chr(10).join(report_data.get('suggestions', []))}

        Best regards,
        Vocalytics Team
        
        ---
        Generated by Vocalytics AI-Powered Call Center Dashboard
        """
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send email with improved error handling
        try:
            server = smtplib.SMTP(EMAIL_HOST, EMAIL_PORT)
            server.starttls()
            server.login(EMAIL_USER, EMAIL_PASSWORD)
            text = msg.as_string()
            server.sendmail(EMAIL_USER, recipient_email, text)
            server.quit()
            
            return jsonify({
                'success': True,
                'message': f'Email sent successfully to {recipient_email}',
                'demo_mode': False
            })
            
        except smtplib.SMTPAuthenticationError as e:
            return jsonify({
                'error': 'Email authentication failed. Please check EMAIL_USER and EMAIL_PASSWORD configuration.',
                'details': 'For Gmail, use App Password instead of regular password.',
                'setup_instructions': 'https://support.google.com/accounts/answer/185833'
            }), 401
            
        except smtplib.SMTPRecipientsRefused as e:
            return jsonify({
                'error': f'Invalid recipient email address: {recipient_email}',
                'details': str(e)
            }), 400
            
        except smtplib.SMTPServerDisconnected as e:
            return jsonify({
                'error': 'Email server connection failed. Please try again.',
                'details': str(e)
            }), 503
            
        except Exception as e:
            return jsonify({
                'error': f'Email sending failed: {str(e)}',
                'suggestion': 'Check your email configuration and network connection.'
            }), 500
        
    except Exception as e:
        print(f"Email sending error: {str(e)}")
        return jsonify({'error': f'Failed to process email request: {str(e)}'}), 500

@app.route('/api/email/status', methods=['GET'])
def get_email_status():
    """Get email configuration status"""
    return jsonify({
        'email_enabled': EMAIL_ENABLED,
        'email_configured': EMAIL_USER is not None,
        'smtp_host': EMAIL_HOST,
        'smtp_port': EMAIL_PORT,
        'demo_mode': not EMAIL_ENABLED,
        'setup_instructions': {
            'step1': 'Set EMAIL_USER environment variable to your email address',
            'step2': 'Set EMAIL_PASSWORD environment variable to your app password',
            'step3': 'For Gmail, generate an App Password: https://support.google.com/accounts/answer/185833',
            'step4': 'Restart the application to apply changes'
        }
    })

@app.route('/api/email/test', methods=['POST'])
def test_email():
    """Test email configuration"""
    try:
        if not EMAIL_ENABLED:
            return jsonify({
                'success': False,
                'message': 'Email not configured. Set EMAIL_USER and EMAIL_PASSWORD environment variables.',
                'demo_mode': True
            })
        
        data = request.get_json()
        test_email = data.get('email', EMAIL_USER)
        
        # Create test message
        msg = MIMEMultipart()
        msg['From'] = EMAIL_USER
        msg['To'] = test_email
        msg['Subject'] = 'Vocalytics - Email Test'
        
        body = """
        This is a test email from Vocalytics AI-Powered Call Center Dashboard.
        
        If you received this email, your email configuration is working correctly!
        
        Best regards,
        Vocalytics Team
        """
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send test email
        server = smtplib.SMTP(EMAIL_HOST, EMAIL_PORT)
        server.starttls()
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        server.sendmail(EMAIL_USER, test_email, msg.as_string())
        server.quit()
        
        return jsonify({
            'success': True,
            'message': f'Test email sent successfully to {test_email}',
            'demo_mode': False
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Email test failed: {str(e)}',
            'suggestion': 'Check your email configuration and credentials.'
        }), 500

# Call Management Routes
@app.route('/api/calls/initiate', methods=['POST'])
def initiate_call():
    """Initiate a new outbound call"""
    try:
        data = request.get_json()
        if not data or 'phone_number' not in data:
            return jsonify({'error': 'Phone number is required'}), 400
        
        phone_number = data['phone_number']
        
        # Validate phone number
        validation = call_service.validate_phone_number(phone_number)
        if not validation['valid']:
            return jsonify({'error': validation['error']}), 400
        
        formatted_number = validation['formatted_number']
        
        # Get webhook base URL (you'll need to set this for production)
        webhook_base_url = request.host_url.rstrip('/')
        
        # Initiate the call
        result = call_service.initiate_call(formatted_number, webhook_base_url)
        
        if result['success']:
            # Emit call status to connected clients
            socketio.emit('call_initiated', {
                'call_uuid': result['call_uuid'],
                'call_sid': result['call_sid'],
                'to_number': formatted_number,
                'status': 'initiated'
            })
            
            return jsonify(result)
        else:
            return jsonify(result), 500
            
    except Exception as e:
        return jsonify({'error': f'Failed to initiate call: {str(e)}'}), 500

@app.route('/api/call/start', methods=['POST'])
def start_call():
    """Start a new call - Frontend compatible endpoint"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Request body is required'}), 400
        
        phone_number = data.get('phone_number')
        if not phone_number:
            return jsonify({'error': 'Phone number is required'}), 400
        
        # Get optional parameters
        enable_transcription = data.get('enable_transcription', True)
        enable_analysis = data.get('enable_analysis', True)
        enable_recording = data.get('enable_recording', True)
        call_type = data.get('call_type', 'outbound')
        
        # Validate phone number
        validation = call_service.validate_phone_number(phone_number)
        if not validation['valid']:
            return jsonify({'error': validation['error']}), 400
        
        formatted_number = validation['formatted_number']
        
        # Get webhook base URL from environment or request
        webhook_base_url = os.getenv('WEBHOOK_BASE_URL') or request.host_url.rstrip('/')
        
        # Initiate the call
        result = call_service.initiate_call(formatted_number, webhook_base_url)
        
        if result['success']:
            # Store call configuration
            call_config = {
                'call_uuid': result['call_uuid'],
                'call_sid': result['call_sid'],
                'phone_number': formatted_number,
                'call_type': call_type,
                'enable_transcription': enable_transcription,
                'enable_analysis': enable_analysis,
                'enable_recording': enable_recording,
                'start_time': datetime.now().isoformat(),
                'status': 'initiated'
            }
            
            # Enable real-time transcription if requested
            if enable_transcription:
                try:
                    call_service.enable_real_time_transcription(result['call_uuid'])
                except Exception as e:
                    print(f"Warning: Failed to enable real-time transcription: {e}")
            
            # Emit call status to connected clients
            socketio.emit('call_started', call_config)
            
            return jsonify({
                'success': True,
                'call_uuid': result['call_uuid'],
                'call_sid': result['call_sid'],
                'phone_number': formatted_number,
                'status': 'initiated',
                'config': call_config,
                'message': 'Call started successfully'
            })
        else:
            return jsonify(result), 500
            
    except Exception as e:
        return jsonify({'error': f'Failed to start call: {str(e)}'}), 500

@app.route('/api/call/webhook', methods=['POST'])
def handle_call_webhook_api():
    """Handle call webhook - Returns TwiML XML for call handling"""
    try:
        from twilio.twiml.voice_response import VoiceResponse
        
        # Get form data from Twilio
        data = request.form.to_dict()
        
        # Extract call information
        call_sid = data.get('CallSid')
        call_status = data.get('CallStatus')
        from_number = data.get('From')
        to_number = data.get('To')
        
        # Log webhook event
        webhook_service.log_webhook_event('call_webhook_api', call_sid, data)
        
        # Create TwiML response
        response = VoiceResponse()
        
        # Greeting message
        response.say(
            "Hello! Thank you for calling. Your call is being processed and recorded for quality assurance.",
            voice='Polly.Joanna',
            language='en-US'
        )
        
        # Start recording with transcription
        webhook_base_url = os.getenv('WEBHOOK_BASE_URL', 'https://9fb47f1115d0.ngrok-free.app')
        
        response.record(
            action=f'{webhook_base_url}/api/call/recording',
            method='POST',
            max_length=3600,  # 1 hour max
            play_beep=False,
            record_on_hangup=True,
            transcribe=True,
            transcribe_callback=f'{webhook_base_url}/api/call/transcription',
            dual_channel=True
        )
        
        # Add a pause for conversation
        response.pause(length=2)
        
        # Final message
        response.say(
            "Thank you for your call. Have a great day!",
            voice='Polly.Joanna',
            language='en-US'
        )
        
        # Emit status update to connected clients
        socketio.emit('call_webhook_received', {
            'call_sid': call_sid,
            'status': call_status,
            'from': from_number,
            'to': to_number,
            'data': data
        })
        
        return str(response), 200, {'Content-Type': 'text/xml'}
        
    except Exception as e:
        print(f"Call webhook error: {str(e)}")
        # Return a simple TwiML response even on error
        from twilio.twiml.voice_response import VoiceResponse
        response = VoiceResponse()
        response.say("Sorry, there was an error processing your call. Please try again later.")
        return str(response), 500, {'Content-Type': 'text/xml'}

@app.route('/api/call/status', methods=['POST'])
def handle_call_status_api():
    """Handle call status updates - Frontend compatible endpoint"""
    try:
        data = request.get_json() or request.form.to_dict()
        
        # Extract call information
        call_sid = data.get('CallSid') or data.get('call_sid')
        call_status = data.get('CallStatus') or data.get('status')
        
        if not call_sid:
            return jsonify({'error': 'CallSid is required'}), 400
        
        # Log status update
        webhook_service.log_webhook_event('call_status_api', call_sid, data)
        
        # Emit status update to connected clients
        socketio.emit('call_status_updated', {
            'call_sid': call_sid,
            'status': call_status,
            'data': data
        })
        
        return jsonify({
            'success': True,
            'call_sid': call_sid,
            'status': call_status,
            'message': 'Status updated successfully'
        })
        
    except Exception as e:
        return jsonify({'error': f'Failed to update status: {str(e)}'}), 500

@app.route('/api/call/recording', methods=['POST'])
def handle_call_recording_api():
    """Handle recording webhook - Returns TwiML XML"""
    try:
        from twilio.twiml.voice_response import VoiceResponse
        
        # Get form data from Twilio
        data = request.form.to_dict()
        
        # Extract recording information
        call_sid = data.get('CallSid')
        recording_url = data.get('RecordingUrl')
        recording_sid = data.get('RecordingSid')
        recording_duration = data.get('RecordingDuration')
        
        # Log recording event
        webhook_service.log_webhook_event('call_recording_api', call_sid, data)
        
        # Process recording if available
        if recording_url and recording_sid:
            # Emit recording info to connected clients
            socketio.emit('call_recording_ready', {
                'call_sid': call_sid,
                'recording_sid': recording_sid,
                'recording_url': recording_url,
                'duration': recording_duration,
                'data': data
            })
        
        # Create TwiML response to continue or end call
        response = VoiceResponse()
        response.say("Thank you for your call. Goodbye!", voice='Polly.Joanna')
        response.hangup()
        
        return str(response), 200, {'Content-Type': 'text/xml'}
        
    except Exception as e:
        print(f"Recording webhook error: {str(e)}")
        from twilio.twiml.voice_response import VoiceResponse
        response = VoiceResponse()
        response.hangup()
        return str(response), 500, {'Content-Type': 'text/xml'}

@app.route('/api/call/transcription', methods=['POST'])
def handle_call_transcription_api():
    """Handle transcription webhook - Returns TwiML XML"""
    try:
        from twilio.twiml.voice_response import VoiceResponse
        
        # Get form data from Twilio
        data = request.form.to_dict()
        
        # Extract transcription information
        call_sid = data.get('CallSid')
        transcription_text = data.get('TranscriptionText')
        transcription_status = data.get('TranscriptionStatus')
        transcription_url = data.get('TranscriptionUrl')
        
        # Log transcription event
        webhook_service.log_webhook_event('call_transcription_api', call_sid, data)
        
        # Process transcription if available
        if transcription_text:
            # Emit transcription to connected clients
            socketio.emit('call_transcription_ready', {
                'call_sid': call_sid,
                'transcription': transcription_text,
                'status': transcription_status,
                'url': transcription_url,
                'data': data
            })
        
        # Return empty TwiML (no action needed)
        response = VoiceResponse()
        return str(response), 200, {'Content-Type': 'text/xml'}
        
    except Exception as e:
        print(f"Transcription webhook error: {str(e)}")
        from twilio.twiml.voice_response import VoiceResponse
        response = VoiceResponse()
        return str(response), 500, {'Content-Type': 'text/xml'}

@app.route('/api/calls/end/<call_sid>', methods=['POST'])
def end_call(call_sid):
    """End an active call"""
    try:
        result = call_service.end_call(call_sid)
        
        if result['success']:
            # Emit call ended status
            socketio.emit('call_ended', {
                'call_sid': call_sid,
                'status': 'completed'
            })
            
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to end call: {str(e)}'}), 500

@app.route('/api/calls/status/<call_sid>', methods=['GET'])
def get_call_status(call_sid):
    """Get current status of a call"""
    try:
        result = call_service.get_call_status(call_sid)
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to get call status: {str(e)}'}), 500

# Twilio Webhook Routes
@app.route('/webhook/call/<call_uuid>', methods=['POST'])
def handle_call_webhook(call_uuid):
    """Handle Twilio call webhook"""
    try:
        print(f"🔥 Call webhook triggered for UUID: {call_uuid}")
        
        # Generate TwiML response for the call
        twiml = call_service.generate_call_twiml(call_uuid)
        
        print(f"🔥 Generated TwiML: {twiml}")
        
        # Update call status
        call_service.update_call_status(call_uuid, 'in_progress')
        
        # Emit status update
        socketio.emit('call_status_update', {
            'call_uuid': call_uuid,
            'status': 'in_progress'
        })
        
        return twiml, 200, {'Content-Type': 'text/xml'}
        
    except Exception as e:
        print(f"Call webhook error: {str(e)}")
        # Return valid TwiML even on error
        from twilio.twiml.voice_response import VoiceResponse
        response = VoiceResponse()
        response.say("Hello! Thank you for calling. Your call is being processed.")
        response.record(max_length=300)
        response.say("Thank you for your call. Goodbye!")
        response.hangup()
        return str(response), 200, {'Content-Type': 'text/xml'}

@app.route('/webhook/status/<call_uuid>', methods=['POST'])
def handle_status_webhook(call_uuid):
    """Handle Twilio call status webhook"""
    try:
        # Validate webhook signature for security (disabled for development)
        signature = request.headers.get('X-Twilio-Signature', '')
        request_url = request.url
        post_data = request.form.to_dict()
        
        # TODO: Re-enable webhook signature validation for production
        # if not webhook_service.validate_webhook_signature(request_url, post_data, signature):
        #     return 'Unauthorized', 401
        
        # Log webhook event
        webhook_service.log_webhook_event('call_status', call_uuid, post_data)
        
        # Handle the webhook using the webhook service
        result = webhook_service.handle_call_status_webhook(call_uuid, post_data)
        
        if result['success']:
            return 'OK', 200
        else:
            return f"Error: {result['error']}", 500
        
    except Exception as e:
        print(f"Status webhook error: {str(e)}")
        return f"Error: {str(e)}", 500

@app.route('/webhook/recording/<call_uuid>', methods=['POST'])
def handle_recording_webhook(call_uuid):
    """Handle Twilio recording webhook"""
    try:
        # Validate webhook signature for security (disabled for development)
        signature = request.headers.get('X-Twilio-Signature', '')
        request_url = request.url
        post_data = request.form.to_dict()
        
        # TODO: Re-enable webhook signature validation for production
        # if not webhook_service.validate_webhook_signature(request_url, post_data, signature):
        #     return 'Unauthorized', 401
        
        # Log webhook event
        webhook_service.log_webhook_event('recording_status', call_uuid, post_data)
        
        # Handle the webhook using the webhook service
        result = webhook_service.handle_recording_status_webhook(call_uuid, post_data)
        
        if result['success']:
            return 'OK', 200
        else:
            return f"Error: {result['error']}", 500
        
    except Exception as e:
        print(f"Recording webhook error: {str(e)}")
        return f"Error: {str(e)}", 500

@app.route('/webhook/transcription/<call_uuid>', methods=['POST'])
def handle_transcription_webhook(call_uuid):
    """Handle Twilio transcription webhook"""
    try:
        # Validate webhook signature for security (disabled for development)
        signature = request.headers.get('X-Twilio-Signature', '')
        request_url = request.url
        post_data = request.form.to_dict()
        
        # TODO: Re-enable webhook signature validation for production
        # if not webhook_service.validate_webhook_signature(request_url, post_data, signature):
        #     return 'Unauthorized', 401
        
        # Log webhook event
        webhook_service.log_webhook_event('transcription', call_uuid, post_data)
        
        # Handle the webhook using the webhook service
        result = webhook_service.handle_transcription_webhook(call_uuid, post_data)
        
        if result['success']:
            return 'OK', 200
        else:
            return f"Error: {result['error']}", 500
        
    except Exception as e:
        print(f"Transcription webhook error: {str(e)}")
        return f"Error: {str(e)}", 500

# Recording Management Routes
@app.route('/api/recordings/list', methods=['GET'])
def list_recordings():
    """List all local recordings"""
    try:
        result = recording_service.get_local_recordings()
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to list recordings: {str(e)}'}), 500

@app.route('/api/recordings/call/<call_sid>', methods=['GET'])
def get_call_recordings(call_sid):
    """Get recordings for a specific call"""
    try:
        result = recording_service.list_call_recordings(call_sid)
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to get call recordings: {str(e)}'}), 500

@app.route('/api/recordings/info/<recording_sid>', methods=['GET'])
def get_recording_info(recording_sid):
    """Get information about a specific recording"""
    try:
        result = recording_service.get_recording_info(recording_sid)
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to get recording info: {str(e)}'}), 500

@app.route('/api/recordings/download/<recording_sid>', methods=['POST'])
def download_recording_manually(recording_sid):
    """Manually download a recording"""
    try:
        data = request.get_json()
        call_uuid = data.get('call_uuid', 'manual')
        call_sid = data.get('call_sid', '')
        
        # Get recording info first
        recording_info = recording_service.get_recording_info(recording_sid)
        
        if not recording_info['success']:
            return jsonify(recording_info), 404
        
        # Download the recording
        media_url = recording_info['media_url']
        result = recording_service.download_recording(media_url, call_uuid, call_sid)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to download recording: {str(e)}'}), 500

@app.route('/api/recordings/delete/<recording_sid>', methods=['DELETE'])
def delete_recording_from_twilio(recording_sid):
    """Delete a recording from Twilio"""
    try:
        result = recording_service.delete_recording(recording_sid)
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to delete recording: {str(e)}'}), 500

@app.route('/api/recordings/cleanup', methods=['POST'])
def cleanup_old_recordings():
    """Clean up old local recordings"""
    try:
        data = request.get_json()
        days_old = data.get('days_old', 30) if data else 30
        
        result = recording_service.cleanup_old_recordings(days_old)
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to cleanup recordings: {str(e)}'}), 500

@app.route('/api/recordings/analyze/<filename>', methods=['POST'])
def analyze_local_recording():
    """Analyze a local recording file"""
    try:
        data = request.get_json()
        filename = data.get('filename')
        
        if not filename:
            return jsonify({'error': 'Filename is required'}), 400
        
        # Check if file exists in recordings folder
        file_path = os.path.join(recording_service.recordings_folder, filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'Recording file not found'}), 404
        
        # Copy file to uploads folder for analysis
        upload_filename = f"recording_{filename}"
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], upload_filename)
        
        import shutil
        shutil.copy2(file_path, upload_path)
        
        # Analyze the recording using existing pipeline
        if whisper_model is None:
            return jsonify({'error': 'Whisper model not available'}), 500
        
        # Convert audio to compatible format if needed
        audio_path = preprocess_audio(upload_path)
        
        # Transcribe using Whisper
        result = whisper_model.transcribe(audio_path, word_timestamps=True)
        
        # Extract transcription and segments
        transcription = result['text']
        segments = []
        
        for segment in result['segments']:
            segment_data = {
                'start': segment['start'],
                'end': segment['end'],
                'text': segment['text'],
                'words': []
            }
            
            if 'words' in segment:
                for word in segment['words']:
                    segment_data['words'].append({
                        'start': word['start'],
                        'end': word['end'],
                        'word': word['word']
                    })
            
            segments.append(segment_data)
        
        # Analyze the transcription
        analysis_result = {
            'sentiment': analyze_sentiment(transcription),
            'emotions': analyze_emotions(transcription),
            'topics': extract_key_topics(transcription),
            'statistics': get_text_statistics(transcription),
            'summary': generate_summary(transcription),
            'insights': extract_insights(transcription),
            'suggestions': generate_suggestions(transcription)
        }
        
        # Clean up temporary files
        if audio_path != upload_path:
            os.remove(audio_path)
        os.remove(upload_path)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'transcription': transcription,
            'segments': segments,
            'language': result.get('language', 'unknown'),
            'analysis': analysis_result
        })
        
    except Exception as e:
        return jsonify({'error': f'Failed to analyze recording: {str(e)}'}), 500

# Webhook Management Routes
@app.route('/api/webhooks/stats', methods=['GET'])
def get_webhook_stats():
    """Get webhook processing statistics"""
    try:
        stats = webhook_service.get_webhook_statistics()
        return jsonify(stats)
        
    except Exception as e:
        return jsonify({'error': f'Failed to get webhook stats: {str(e)}'}), 500

@app.route('/api/webhooks/test', methods=['POST'])
def test_webhook():
    """Test webhook connectivity"""
    try:
        data = request.get_json()
        test_type = data.get('type', 'call_status')
        call_uuid = data.get('call_uuid', 'test-uuid')
        
        # Create test webhook data
        test_data = {
            'CallSid': 'test-call-sid',
            'CallStatus': 'completed',
            'CallDuration': '30',
            'From': '+1234567890',
            'To': '+0987654321'
        }
        
        if test_type == 'call_status':
            result = webhook_service.handle_call_status_webhook(call_uuid, test_data)
        elif test_type == 'recording':
            test_data.update({
                'RecordingStatus': 'completed',
                'RecordingSid': 'test-recording-sid',
                'RecordingUrl': 'https://api.twilio.com/test-recording',
                'RecordingDuration': '25'
            })
            result = webhook_service.handle_recording_status_webhook(call_uuid, test_data)
        elif test_type == 'transcription':
            test_data.update({
                'TranscriptionText': 'This is a test transcription',
                'TranscriptionStatus': 'completed',
                'TranscriptionSid': 'test-transcription-sid',
                'RecordingSid': 'test-recording-sid'
            })
            result = webhook_service.handle_transcription_webhook(call_uuid, test_data)
        else:
            return jsonify({'error': 'Invalid test type'}), 400
        
        return jsonify({
            'success': True,
            'test_type': test_type,
            'result': result
        })
        
    except Exception as e:
        return jsonify({'error': f'Webhook test failed: {str(e)}'}), 500

@app.route('/api/webhooks/logs', methods=['GET'])
def get_webhook_logs():
    """Get recent webhook logs"""
    try:
        limit = request.args.get('limit', 50, type=int)
        logs = []
        
        # Read recent logs from file
        log_file = 'webhook_events.log'
        if os.path.exists(log_file):
            with open(log_file, 'r') as f:
                lines = f.readlines()
                for line in lines[-limit:]:
                    try:
                        log_entry = json.loads(line.strip())
                        logs.append(log_entry)
                    except json.JSONDecodeError:
                        continue
        
        return jsonify({
            'success': True,
            'logs': logs,
            'count': len(logs)
        })
        
    except Exception as e:
        return jsonify({'error': f'Failed to get webhook logs: {str(e)}'}), 500

# Enhanced Machine Detection Webhook
@app.route('/webhook/machine-detection/<call_uuid>', methods=['POST'])
def handle_machine_detection_webhook(call_uuid):
    """Handle machine detection webhook"""
    try:
        # Validate webhook signature
        signature = request.headers.get('X-Twilio-Signature', '')
        request_url = request.url
        post_data = request.form.to_dict()
        
        if not webhook_service.validate_webhook_signature(request_url, post_data, signature):
            return 'Unauthorized', 401
        
        answered_by = request.form.get('AnsweredBy')
        call_sid = request.form.get('CallSid')
        machine_detection_duration = request.form.get('MachineDetectionDuration')
        
        # Log the detection result
        webhook_service.log_webhook_event('machine_detection', call_uuid, post_data)
        
        # Update call info
        call_service.update_call_status(
            call_uuid,
            'machine_detected',
            answered_by=answered_by,
            machine_detection_duration=machine_detection_duration
        )
        
        # Emit machine detection result
        socketio.emit('machine_detection', {
            'call_uuid': call_uuid,
            'call_sid': call_sid,
            'answered_by': answered_by,
            'detection_duration': machine_detection_duration,
            'timestamp': datetime.now().isoformat()
        })
        
        # If it's a machine, you might want to leave a voicemail or hang up
        if answered_by in ['machine_start', 'machine_end_beep', 'machine_end_silence']:
            # Generate TwiML for machine handling
            from twilio.twiml.voice_response import VoiceResponse
            response = VoiceResponse()
            response.say(
                "Hello, this is an automated call. Please call us back at your convenience.",
                voice='Polly.Joanna',
                language='en-US'
            )
            response.hangup()
            return str(response), 200, {'Content-Type': 'text/xml'}
        
        return 'OK', 200
        
    except Exception as e:
        print(f"Machine detection webhook error: {str(e)}")
        return f"Error: {str(e)}", 500

# WebSocket Events
@socketio.on('connect')
def handle_connect():
    """Handle client connection"""
    print('Client connected')
    emit('connected', {'message': 'Connected to call service'})

@socketio.on('disconnect')
def handle_disconnect():
    """Handle client disconnection"""
    print('Client disconnected')

@socketio.on('request_call_status')
def handle_call_status_request(data):
    """Handle request for call status"""
    try:
        call_uuid = data.get('call_uuid')
        if call_uuid:
            call_info = call_service.get_call_info(call_uuid)
            if call_info:
                emit('call_status_response', {
                    'call_uuid': call_uuid,
                    'call_info': call_info
                })
            else:
                emit('call_status_response', {
                    'call_uuid': call_uuid,
                    'error': 'Call not found'
                })
    except Exception as e:
        emit('call_status_response', {
            'error': str(e)
        })

@app.route('/test/aws-transcribe')
def aws_transcribe_test():
    """Serve AWS Transcribe Streaming test page"""
    return render_template('aws_transcribe_test.html')

@app.route('/static/js/<path:filename>')
def serve_js(filename):
    """Serve JavaScript files"""
    return send_from_directory('../js', filename)

@app.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy', 
        'timestamp': datetime.now().isoformat(),
        'websocket_server': aws_transcribe_service.websocket_server_started,
        'aws_available': aws_transcribe_service.is_aws_available(),
        'websocket_port': aws_transcribe_service.server_port
    })

@app.route('/api/aws/transcribe/check', methods=['GET'])
def check_aws_transcribe():
    """Check if AWS Transcribe streaming is available"""
    try:
        return jsonify({
            'available': aws_transcribe_service.is_aws_available(),
            'message': 'AWS Transcribe streaming ready' if aws_transcribe_service.is_aws_available() else 'AWS Transcribe not configured',
            'websocket_url': f'ws://localhost:{aws_transcribe_service.server_port}/aws-transcribe'
        })
    except Exception as e:
        return jsonify({
            'available': False,
            'error': str(e)
        }), 500

@app.route('/api/aws/transcribe/start', methods=['POST'])
def start_aws_transcribe_session():
    """Start AWS Transcribe streaming session"""
    try:
        data = request.get_json()
        call_uuid = data.get('call_uuid')
        language = data.get('language', 'en-US')
        enable_diarization = data.get('enable_diarization', True)
        max_speakers = data.get('max_speakers', 4)
        
        if not call_uuid:
            return jsonify({'error': 'call_uuid is required'}), 400
        
        # Generate session configuration
        session_config = {
            'call_uuid': call_uuid,
            'language': language,
            'enable_diarization': enable_diarization,
            'max_speakers': max_speakers,
            'websocket_url': f'ws://localhost:{aws_transcribe_service.server_port}/aws-transcribe'
        }
        
        return jsonify({
            'success': True,
            'config': session_config,
            'aws_available': aws_transcribe_service.is_aws_available()
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/aws/transcribe/sessions', methods=['GET'])
def get_transcribe_sessions():
    """Get active transcription sessions"""
    try:
        sessions = aws_transcribe_service.get_active_sessions()
        return jsonify({
            'success': True,
            'sessions': sessions,
            'count': len(sessions)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/aws/transcribe/sessions/<session_id>', methods=['GET'])
def get_transcribe_session(session_id):
    """Get specific transcription session"""
    try:
        session_data = aws_transcribe_service.get_session_transcript(session_id)
        if not session_data:
            return jsonify({'error': 'Session not found'}), 404
        
        return jsonify({
            'success': True,
            'session': session_data
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/aws/transcribe/sessions/<session_id>/stop', methods=['POST'])
def stop_transcribe_session(session_id):
    """Stop transcription session"""
    try:
        # The actual stopping is handled via WebSocket
        # This endpoint is for REST API compatibility
        return jsonify({
            'success': True,
            'message': 'Stop signal sent',
            'session_id': session_id
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Add real-time analysis handler for AWS Transcribe
@socketio.on('real_time_analysis_request')
def handle_real_time_analysis(data):
    """Handle real-time analysis requests from AWS Transcribe"""
    try:
        call_uuid = data.get('call_uuid')
        session_id = data.get('session_id')
        text = data.get('text', '')
        current_utterance = data.get('current_utterance', {})
        speaker_info = data.get('speaker_info', {})
        
        if not text.strip():
            return
        
        # Perform quick analysis
        analysis_result = {
            'sentiment': analyze_sentiment(text),
            'emotions': analyze_emotions(text),
            'topics': extract_key_topics(text),
            'statistics': get_text_statistics(text)
        }
        
        # Emit analysis results
        emit('real_time_analysis_result', {
            'call_uuid': call_uuid,
            'session_id': session_id,
            'text': text,
            'analysis': analysis_result,
            'speaker': speaker_info,
            'utterance': current_utterance,
            'timestamp': datetime.now().isoformat()
        })
        
        # Also emit to general call analysis
        emit('live_analysis_update', {
            'call_uuid': call_uuid,
            'analysis': analysis_result,
            'speaker': speaker_info.get('name', 'Unknown'),
            'text_sample': text[-100:] if len(text) > 100 else text,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"Real-time analysis error: {str(e)}")
        emit('analysis_error', {
            'call_uuid': data.get('call_uuid'),
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        })

if __name__ == '__main__':
    # Start WebSocket server after Flask app is ready
    aws_transcribe_service.start_websocket_server()
    socketio.run(app, debug=True, host='0.0.0.0', port=5000) 